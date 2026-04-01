# Export utilities for generating DOC and PDF documents
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, KeepTogether
from reportlab.lib.enums import TA_LEFT, TA_RIGHT
from reportlab.lib import colors
from io import BytesIO
import re

def escape_html_tags(text):
    """Remove or escape HTML-like tags from text"""
    if not text:
        return ""
    # Remove <st> tags but keep content
    text = re.sub(r'<st>(.*?)</st>', r'\1', text)
    # Remove any other HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    return text.strip()

def generate_doc(questions, filename="questions_export.docx", sort_by='id', db_name='database'):
    """Generate a DOCX document with questions"""
    # Sort questions based on sort_by parameter
    if sort_by == 'category':
        # Sort by primary_domain, then by subdomain, then by id
        sorted_questions = sorted(questions, key=lambda q: (
            q.get('primary_domain', 'zzz'),  # zzz to put 'indefinito' at end
            q.get('subdomain', 'zzz'),
            q.get('id', '')
        ))
    else:  # sort_by == 'id'
        # Sort by ID (extract number and sort numerically)
        def extract_number(q_id):
            import re
            match = re.search(r'(\d+)', q_id)
            return int(match.group(1)) if match else 0
        sorted_questions = sorted(questions, key=lambda q: extract_number(q.get('id', '')))
    
    doc = Document()
    
    # Title - use database name
    title = doc.add_heading(f'Lista domande - {db_name}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add summary
    doc.add_paragraph(f'Totale domande esportate: {len(questions)}')
    doc.add_paragraph('')
    
    # Add questions
    for idx, question in enumerate(sorted_questions, 1):
        q_id = question.get('id', f'Q{idx}')
        raw_text = escape_html_tags(question.get('raw_text', 'Nessun testo'))
        answers = question.get('answers', {})
        correct = question.get('correct', [])
        primary_domain = question.get('primary_domain', '')
        subdomain = question.get('subdomain', '')
        
        # Get correct answers string
        correct_letters = [c for c in correct if c and c != 'null']
        correct_str = f'[{", ".join(sorted(correct_letters))}]' if correct_letters else ''
        
        # Add category info paragraph (small text above the question)
        category_info = f'{primary_domain} / {subdomain}'
        p_category = doc.add_paragraph()
        p_category.paragraph_format.space_after = Pt(2)  # Small space after category info
        p_category.paragraph_format.space_before = Pt(0)
        run_category = p_category.add_run(category_info)
        run_category.font.size = Pt(7)  # Even smaller font size
        run_category.font.italic = True  # Italic style for category info
        
        # Question with ID - use hanging indent for alignment
        # Calculate indent dynamically based on ID length
        id_width = len(q_id) * 7  # Approximate width in points (7pt per character)
        
        p_question = doc.add_paragraph()
        p_question.paragraph_format.space_after = Pt(4)
        p_question.paragraph_format.line_spacing = 1.0
        # Keep question with next paragraph (answers)
        p_question.paragraph_format.keep_with_next = True
        # Hanging indent: first line at 0, subsequent lines indented
        p_question.paragraph_format.first_line_indent = Pt(-id_width)  # Negative = hanging
        p_question.paragraph_format.left_indent = Pt(id_width)  # Where subsequent lines start
        
        run_id = p_question.add_run(f'{q_id} ')
        run_id.font.size = Pt(10)
        run_id.bold = True
        
        run_text = p_question.add_run(raw_text)
        run_text.font.size = Pt(10)
        
        # Add correct answers on the right
        if correct_str:
            run_correct = p_question.add_run(f'  {correct_str}')
            run_correct.font.size = Pt(10)
            run_correct.bold = True
        
        # Answers - compact, no highlighting, aligned with question text
        answer_letters = sorted(answers.keys())
        for i, letter in enumerate(answer_letters):
            answer_text = escape_html_tags(answers.get(letter, ''))
            if not answer_text:
                continue
            
            # Calculate answer indent: ID width + letter width
            letter_width = 14  # "A) " is approximately 14 points
            
            p_answer = doc.add_paragraph()
            p_answer.paragraph_format.space_before = Pt(1)
            p_answer.paragraph_format.space_after = Pt(1)
            p_answer.paragraph_format.line_spacing = 1.0
            # Keep all answers together with next (except last one)
            if i < len(answer_letters) - 1:
                p_answer.paragraph_format.keep_with_next = True
            # Hanging indent for answers: letter aligned, text wraps to same position
            p_answer.paragraph_format.first_line_indent = Pt(-letter_width)  # Negative for letter
            p_answer.paragraph_format.left_indent = Pt(id_width + letter_width)  # Align with question text
            
            # Normal answer
            run_letter = p_answer.add_run(f'{letter}) ')
            run_letter.font.size = Pt(10)
            run_answer = p_answer.add_run(answer_text)
            run_answer.font.size = Pt(10)
        
        # Add subtle separator
        if idx < len(questions):
            p_sep = doc.add_paragraph()
            p_sep.paragraph_format.space_before = Pt(6)
            p_sep.paragraph_format.space_after = Pt(6)
            run_sep = p_sep.add_run('· · ·')
            run_sep.font.color.rgb = RGBColor(200, 200, 20)
            run_sep.font.size = Pt(8)
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer, filename

def generate_pdf(questions, filename="questions_export.pdf", sort_by='id', db_name='database'):
    """Generate a PDF document with questions"""
    # Sort questions based on sort_by parameter
    if sort_by == 'category':
        # Sort by primary_domain, then by subdomain, then by id
        sorted_questions = sorted(questions, key=lambda q: (
            q.get('primary_domain', 'zzz'),  # zzz to put 'indefinito' at end
            q.get('subdomain', 'zzz'),
            q.get('id', '')
        ))
    else:  # sort_by == 'id'
        # Sort by ID (extract number and sort numerically)
        def extract_number(q_id):
            import re
            match = re.search(r'(\d+)', q_id)
            return int(match.group(1)) if match else 0
        sorted_questions = sorted(questions, key=lambda q: extract_number(q.get('id', '')))
    
    buffer = BytesIO()
    
    # Create PDF document
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2*cm,
        bottomMargin=2*cm
    )
    
    # Get styles
    styles = getSampleStyleSheet()
    
    # Custom styles - compact
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=14,
        alignment=1,  # Center
        spaceAfter=15
    )
    
    question_style = ParagraphStyle(
        'QuestionText',
        parent=styles['Normal'],
        fontSize=10,
        spaceAfter=0,
        spaceBefore=0,
        leading=12  # Line spacing
    )
    
    answer_style = ParagraphStyle(
        'AnswerText',
        parent=styles['Normal'],
        fontSize=9,
        spaceAfter=0,
        spaceBefore=0,
        leading=11  # Line spacing
    )
    
    correct_answer_style = ParagraphStyle(
        'CorrectAnswer',
        parent=styles['Normal'],
        fontSize=9,
        leftIndent=15,
        spaceAfter=1,
        spaceBefore=1,
        textColor='green'
    )
    
    separator_style = ParagraphStyle(
        'Separator',
        parent=styles['Normal'],
        fontSize=6,
        alignment=1,  # Center
        spaceBefore=4,
        spaceAfter=4,
        textColor='lightgrey'
    )
    
    # Build content
    story = []
    
    # Title - use database name
    story.append(Paragraph(f'Lista domande - {db_name}', title_style))
    story.append(Paragraph(f'Totale domande esportate: {len(sorted_questions)}', styles['Normal']))
    story.append(Spacer(1, 15))
    
    # Add questions
    for idx, question in enumerate(sorted_questions, 1):
        q_id = question.get('id', f'Q{idx}')
        raw_text = escape_html_tags(question.get('raw_text', 'Nessun testo'))
        answers = question.get('answers', {})
        correct = question.get('correct', [])
        primary_domain = question.get('primary_domain', '')
        subdomain = question.get('subdomain', '')
        
        # Get correct answers string
        correct_letters = [c for c in correct if c and c != 'null']
        correct_str = f'[{", ".join(sorted(correct_letters))}]' if correct_letters else ''
        
        # Create category info paragraph (small text above the question)
        category_info = f'{primary_domain} / {subdomain}'
        category_para = Paragraph(f'<i>{category_info}</i>', ParagraphStyle(
            'CategoryInfo',
            parent=styles['Normal'],
            fontSize=7,  # Even smaller font size
            spaceAfter=0,  # Small space after category info
            spaceBefore=0,
            leftIndent=0,
            rightIndent=0
        ))
        
        # Create a table with two columns: [ID] [Text] + [Correct]
        id_para = Paragraph(f'<b>{q_id}</b>', question_style)
        text_para = Paragraph(raw_text, question_style)
        
        # Correct answers on the right
        correct_para = Paragraph(f'<b>{correct_str}</b>', ParagraphStyle(
            'CorrectRight',
            parent=styles['Normal'],
            fontSize=10,
            alignment=TA_RIGHT
        )) if correct_str else Paragraph('', styles['Normal'])
        
        # Create table with question ID, text, and correct answers
        table_data = [[id_para, text_para, correct_para]]
        table = Table(table_data, colWidths=[doc.width * 0.08, doc.width * 0.77, doc.width * 0.15])
        table.setStyle(TableStyle([
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('ALIGN', (0, 0), (0, 0), 'LEFT'),
            ('ALIGN', (1, 0), (1, 0), 'RIGHT'),
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
            ('TOPPADDING', (0, 0), (-1, -1), 0),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 0),
        ]))
        
        # Build question block with answers
        question_block = [category_para, table]  # Add category info before the table
        
        # Add small space between question and answers
        question_block.append(Spacer(1, 4))
        
        # Answers - compact, no highlighting
        answer_letters = sorted(answers.keys())
        for letter in answer_letters:
            answer_text = escape_html_tags(answers.get(letter, ''))
            if not answer_text:
                continue
            
            # Create answer with hanging indent for proper text wrapping
            answer_para = Paragraph(f'{letter}) {answer_text}', ParagraphStyle(
                'AnswerIndent',
                parent=answer_style,
                leftIndent=doc.width * 0.08 + 12,  # Align with question text + letter width
                firstLineIndent=-12  # Negative to create hanging indent
            ))
            question_block.append(answer_para)
        
        # Wrap question + answers in KeepTogether to prevent page breaks
        story.append(KeepTogether(question_block))
        
        # Add subtle separator
        if idx < len(questions):
            story.append(Paragraph('· · ·', separator_style))
    
    # Build PDF
    doc.build(story)
    buffer.seek(0)
    
    return buffer, filename
